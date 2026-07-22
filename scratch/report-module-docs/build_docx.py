from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path.cwd()
OUT = ROOT / "outputs" / "report-module-docs"
ASSETS = OUT / "assets"
DOCX = OUT / "Melann_Report_Module_Documentation.docx"
ASSETS.mkdir(parents=True, exist_ok=True)

BLUE = "1E3A8A"
DARK = "0F172A"
SLATE = "475569"
GREEN = "16A34A"
RED = "DC2626"
GOLD = "B7791F"
LIGHT = "F8FAFC"
BORDER = "CBD5E1"


def font(size=24, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_report_flow():
    img = Image.new("RGB", (1400, 680), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1400, 92], fill=f"#{BLUE}")
    d.text((44, 27), "Report Module Flow", font=font(34, True), fill="white")
    d.text((44, 103), "From transaction records to screen, print, and exports", font=font(22), fill=f"#{SLATE}")
    boxes = [
        ("Transactions", "Loans, payments,\nexpenses, cash, monitoring", 70, 215, "DB"),
        ("Backend Queries", "Express routes compute\nfiltered report data", 390, 215, "API"),
        ("Report UI", "React report page\nsummaries and details", 710, 215, "UI"),
        ("Output", "Print, PDF via browser,\nCSV/Excel where supported", 1030, 215, "OUT"),
    ]
    for i, (title, body, x, y, tag) in enumerate(boxes):
        rounded_rect(d, [x, y, x + 250, y + 250], 26, "#F8FAFC", f"#{BORDER}", 3)
        d.ellipse([x + 22, y + 24, x + 78, y + 80], fill=f"#{BLUE}")
        d.text((x + 39, y + 39), tag, font=font(14, True), fill="white", anchor="mm")
        d.text((x + 28, y + 104), title, font=font(28, True), fill=f"#{DARK}")
        d.multiline_text((x + 28, y + 150), body, font=font(21), fill=f"#{SLATE}", spacing=8)
        if i < len(boxes) - 1:
            x2 = x + 276
            y2 = y + 125
            d.line([x + 250, y2, x2 + 74, y2], fill=f"#{BLUE}", width=6)
            d.polygon([(x2 + 74, y2), (x2 + 50, y2 - 14), (x2 + 50, y2 + 14)], fill=f"#{BLUE}")
    d.text((70, 535), "Key control point: totals must be computed consistently by the backend, then shown identically in screen, print, and export.", font=font(22, True), fill=f"#{DARK}")
    d.text((70, 578), "Current risk: some reports are screen/print oriented, while DCR still has formula and branch-filtering gaps documented in the PRD.", font=font(21), fill=f"#{SLATE}")
    path = ASSETS / "report-flow.png"
    img.save(path)
    return path


def make_dcr_formula():
    img = Image.new("RGB", (1400, 760), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1400, 92], fill=f"#{DARK}")
    d.text((44, 27), "Daily Cash Report: Cash Position Logic", font=font(34, True), fill="white")
    cols = [
        ("Cash In", ["Beginning cash", "Collections", "Adjustments", "Withdrawals from bank"], GREEN),
        ("Cash Out", ["Loan releases", "Expenses", "Deposits to bank"], RED),
        ("Bank Movement", ["Beginning bank", "Deposits + interest", "Less withdrawals + charges"], BLUE),
    ]
    for idx, (title, items, color) in enumerate(cols):
        x = 70 + idx * 430
        y = 145
        rounded_rect(d, [x, y, x + 360, y + 360], 22, "#F8FAFC", f"#{color}", 3)
        d.text((x + 28, y + 28), title, font=font(28, True), fill=f"#{color}")
        cy = y + 93
        for item in items:
            d.ellipse([x + 30, cy + 7, x + 46, cy + 23], fill=f"#{color}")
            d.text((x + 60, cy), item, font=font(23), fill=f"#{DARK}")
            cy += 66
    rounded_rect(d, [210, 565, 1190, 690], 24, "#EFF6FF", f"#{BLUE}", 3)
    d.text((245, 590), "Ending cash on hand = beginning cash + inflows - releases - expenses - deposits", font=font(24, True), fill=f"#{DARK}")
    d.text((245, 635), "Total cash position = ending cash on hand + ending cash on bank. Variance compares actual cash count vs expected cash.", font=font(21), fill=f"#{SLATE}")
    path = ASSETS / "dcr-formula.png"
    img.save(path)
    return path


def make_report_matrix():
    img = Image.new("RGB", (1400, 860), "white")
    d = ImageDraw.Draw(img)
    d.rectangle([0, 0, 1400, 92], fill=f"#{BLUE}")
    d.text((44, 27), "Report Types at a Glance", font=font(34, True), fill="white")
    reports = [
        ("Collection", "Payments by date range", GREEN),
        ("Releases", "Loans released by date/cycle", BLUE),
        ("Maturity", "Loans by maturity range", GOLD),
        ("Reversed", "Payment reversal audit", RED),
        ("Full Paid", "Completed accounts", GREEN),
        ("Collection Sheet", "Per-collector field list", BLUE),
        ("Disclosure", "Client loan disclosure", DARK),
        ("Monitoring", "Alerts, PTP, resolutions", RED),
        ("DCR", "Daily cash position", GOLD),
    ]
    for i, (name, desc, color) in enumerate(reports):
        row = i // 3
        col = i % 3
        x = 70 + col * 430
        y = 145 + row * 210
        rounded_rect(d, [x, y, x + 360, y + 150], 20, "#F8FAFC", f"#{BORDER}", 2)
        d.rectangle([x, y, x + 360, y + 14], fill=f"#{color}")
        d.text((x + 26, y + 40), name, font=font(26, True), fill=f"#{DARK}")
        d.text((x + 26, y + 84), desc, font=font(21), fill=f"#{SLATE}")
    d.text((70, 795), "Note: Loan Type and Payments Encoded still exist in backend/README references, but are not visible in the current Reports sidebar.", font=font(22, True), fill=f"#{RED}")
    path = ASSETS / "report-types-matrix.png"
    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 17, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 11.5, "1F4D78")]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)


REPORTS = [
    ("Collection Report", "Daily and monthly collection view showing active payments by date range, grouped by collector with summary and drilldown detail.", "Fast cash collection review; supports daily and monthly cycle views; useful for collector accountability.", "Relies on active payment status and correct collector assignment; browser print is the main PDF path; branch filtering is not clearly enforced in all current report endpoints."),
    ("Releases Report", "Daily and monthly loan-release report showing released loans, principal totals, collectors, and loan type split such as new, re-loan, or recon.", "Good for release monitoring, cash-out tracking, and collector production review.", "Some totals use principal while DCR should prefer net proceeds for cash outflow; current UI name says Releases but README still calls this Monthly Releases."),
    ("Loans Maturity Checker / Past Due", "Lists loans by maturity date range and highlights accounts with remaining balance that are active or past due.", "Helps collectors and managers focus on accounts maturing or already overdue.", "Name differs across docs: Past Due Report vs Loans Maturity Checker; days overdue calculation depends on maturity date quality and current date."),
    ("Payments Reversed", "Audit report for payments marked reversed, including amount, customer, loan, collector, reversed user, and reason.", "Supports management review of exceptions and helps detect improper reversals.", "Only as reliable as reversal reason and reversed_at encoding; does not replace a formal approval workflow by itself."),
    ("Full Paid Loans", "Completed loan-account report based on loans with fullpaid status and optional date range.", "Useful for re-loan targeting, completion metrics, and customer lifecycle review.", "Depends on consistent loan status updates; may miss old legacy statuses if mapping is incomplete."),
    ("Collection Sheet", "Printable per-collector active loan list with amortization, balance, maturity, and collected-today fields.", "Strong field-operations tool; provides a ready list for collectors and signatures.", "Requires collector selection; page density can be high; not a management summary by default."),
    ("Disclosure Statement", "Client-specific loan disclosure generated by search or selected loan, with borrower details, loan terms, and schedule.", "Improves transparency and standardizes client-facing loan information.", "Requires accurate customer profile, loan terms, and amortization schedule; search must choose the correct loan when several exist."),
    ("Monitoring Summary", "Operational monitoring report for alerts, escalated accounts, promises to pay, follow-up logs, and resolutions.", "Gives management a high-level view of collection-risk activity beyond ordinary payment reports.", "Backend comment says 10 reports required, but current endpoint returns summary metrics; full report definitions may still need product decisions."),
    ("Daily Cash Report", "Separate Finance report summarizing selected-date releases, expenses, collections, bank movement, cash on hand, and cash position.", "Best daily closing and accounting reconciliation report; supports print, CSV/Excel export, and compliance handoff checklist.", "PRD identifies unresolved gaps: branch filtering, bank formulas, server-side close totals, DCR numbering, and exact legacy parity."),
    ("Loan Type Summary", "Backend/README-supported report concept for breaking down releases or portfolio by loan type/status.", "Useful for product mix and risk segmentation.", "Not currently visible in the Reports sidebar; needs UI decision before presenting as a primary report type."),
    ("Payments Encoded", "Backend/README-supported report for payments encoded within a date range and encoded-by user.", "Useful for teller/accounting productivity and audit reconciliation.", "Not currently visible in the Reports sidebar; overlaps with Collection Report unless positioned as audit-specific."),
]


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for i, width in enumerate(widths):
            table.columns[i].width = Inches(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string(DARK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def main():
    flow = make_report_flow()
    formula = make_dcr_formula()
    matrix = make_report_matrix()

    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Report Module Documentation")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Melann Lending System V2 Modernization | Report Types, Advantages, Problems, and Limitations")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Prepared: July 14, 2026 | Scope: current React/Node/SQLite modernization repo")

    if (ROOT / "client" / "src" / "assets" / "logo.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ROOT / "client" / "src" / "assets" / "logo.png"), width=Inches(1.0))

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Ang Report Module ang pangunahing lugar para makita ng operations, accounting, at management ang collections, releases, loan maturity, reversals, fully paid accounts, disclosure statements, monitoring activity, at daily cash position. Sa modern system, ang report data ay nanggagaling sa SQLite tables sa backend routes, habang ang React UI ang nagpapakita ng summary, drilldown, print view, at export behavior."
    )
    add_bullets(doc, [
        "Current UI-supported report types: Collection, Releases, Loans Maturity Checker, Payments Reversed, Full Paid Loans, Collection Sheet, Disclosure Statement, Monitoring Summary.",
        "Separate Finance report: Daily Cash Report (DCR), dahil may sariling page, formulas, close workflow, and compliance handoff.",
        "Backend/README-supported but not visible in current Reports sidebar: Loan Type Summary and Payments Encoded.",
        "Main limitation: some report names and availability differ between README, backend routes, and current UI, so report inventory should be normalized before final production rollout."
    ])

    doc.add_picture(str(flow), width=Inches(6.7))

    doc.add_heading("2. Report Module Purpose", level=1)
    doc.add_paragraph(
        "Purpose ng module na gawing mabilis, auditable, at print-ready ang operational reporting. Hindi lang ito listahan ng records: ginagamit ito para mag-reconcile ng collections, i-monitor ang loan releases, makita ang risky or overdue accounts, at suportahan ang daily closing."
    )
    add_table(doc, ["Area", "Description"], [
        ("Operations", "Collector performance, collection sheets, maturity monitoring, and release production."),
        ("Accounting", "Collection totals, reversed payments, DCR cash on hand/bank reconciliation, and expense visibility."),
        ("Management", "Exception review, monitoring summary, full paid accounts, and audit trail support."),
        ("Compliance / Client-facing", "Disclosure statement and DCR loan-release checklist handoff to government-compliance workflows."),
    ], [1.65, 4.85])

    doc.add_heading("3. Types of Reports", level=1)
    doc.add_picture(str(matrix), width=Inches(6.7))
    doc.add_paragraph(
        "Below are the report types with practical usage notes. The first nine are the current user-facing or finance-facing reports. The last two are documented/backend-supported candidates that need a UI decision."
    )
    for idx, (name, desc, advantage, limitation) in enumerate(REPORTS, start=1):
        doc.add_heading(f"3.{idx} {name}", level=2)
        add_table(doc, ["Item", "Details"], [
            ("Description", desc),
            ("Advantage", advantage),
            ("Problems / Limitations", limitation),
        ], [1.55, 4.95])

    doc.add_heading("4. Daily Cash Report Focus", level=1)
    doc.add_paragraph(
        "Ang DCR ay mas accounting-critical kaysa ordinary list report dahil may formulas, cash-on-hand/cash-on-bank logic, day closing, variance, and transaction tagging. Dapat backend ang source of truth para hindi magkaiba ang screen, print, CSV/Excel, at persisted close totals."
    )
    doc.add_picture(str(formula), width=Inches(6.7))
    add_table(doc, ["DCR Section", "Purpose"], [
        ("Loan Releases", "Cash-out review for selected date, including customer, collector, loan type, and amount."),
        ("Expenses", "Daily operating cash outflows."),
        ("Adjustments", "Manual cash corrections that should be auditable."),
        ("Collections by Collector", "Cash inflows from payments, passbooks, penalties, and related collector amounts."),
        ("Withdrawal / Deposit", "Transfers between bank and cash on hand."),
        ("Bank Charges / Interest", "Cash-on-bank movements that affect ending bank balance."),
        ("Cash Summary", "Beginning cash, inflows, outflows, expected ending cash, ending bank, and total cash position."),
        ("Signatures", "Prepared, checked, and approved accountability area for print output."),
    ], [1.8, 4.7])

    doc.add_heading("5. Cross-Cutting Advantages", level=1)
    add_bullets(doc, [
        "Centralized report access for users with report permission across Admin, Manager, Teller, and Accounting roles.",
        "Print-friendly browser output hides navigation and controls for cleaner hard copy.",
        "Collector-level grouping helps accountability and day-to-day branch operations.",
        "DCR adds accounting reconciliation, not just transaction listing.",
        "Monitoring Summary connects collection risk activity to follow-up and promise-to-pay data.",
    ])

    doc.add_heading("6. Cross-Cutting Problems and Limitations", level=1)
    add_bullets(doc, [
        "Report inventory mismatch: README says 9 types, Reports UI shows 8 sidebar types, DCR is separate, while backend exposes additional endpoints.",
        "Some reports depend on status mapping from legacy data; inconsistent statuses can hide or overcount accounts.",
        "Browser print/PDF is practical, but not the same as a controlled server-generated PDF renderer.",
        "Branch filtering is a known risk, especially for DCR and reports that join loans/payments/expenses.",
        "DCR PRD identifies current gaps around bank formulas, backend recomputation on close, unique branch/date close rules, and exact legacy report parity.",
        "Very dense reports such as Collection Sheet can become hard to read unless page breaks and print layout are carefully tested.",
    ])

    doc.add_heading("7. Recommended Next Steps", level=1)
    add_table(doc, ["Priority", "Recommendation", "Reason"], [
        ("High", "Normalize report inventory and labels across README, UI, and backend routes.", "Prevents user confusion and missing/duplicate report expectations."),
        ("High", "Make backend formula services the source of truth for DCR and reusable totals.", "Ensures screen, print, export, and close records match."),
        ("High", "Add branch-aware filters and tests for reports that affect cash or collector accountability.", "Avoids cross-branch totals and audit disputes."),
        ("Medium", "Decide whether Loan Type Summary and Payments Encoded should return to the UI.", "They exist in docs/backend but are not current sidebar items."),
        ("Medium", "Add controlled PDF/export generation for critical reports.", "Improves consistency beyond browser print behavior."),
    ], [0.8, 3.1, 2.6])

    doc.add_heading("8. Source Basis", level=1)
    doc.add_paragraph("This document is based on repository inspection of:")
    add_bullets(doc, [
        "README.md report inventory and role permissions.",
        "CODEBASE_PRD.md legacy reporting inventory.",
        "DCR_MODULE_PRD.md DCR behavior, formulas, and known gaps.",
        "server/src/routes/reports.js and server/src/routes/dcr.js API behavior.",
        "client/src/pages/Reports.jsx and client/src/pages/DailyCashReport.jsx current UI behavior.",
    ])

    doc.core_properties.title = "Report Module Documentation"
    doc.core_properties.subject = "Melann Lending System V2 Report Module and Types"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
